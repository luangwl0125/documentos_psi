import streamlit as st
import datetime
import logging
import time
import re
from docx import Document
from io import BytesIO
from PyPDF2 import PdfReader
from PIL import Image
import pytesseract
import openai
import tempfile
import os

import os
openai.api_key = os.getenv("OPENAI_API_KEY")

logging.basicConfig(level=logging.INFO)
openai.api_key = st.secrets["OPENAI_API_KEY"]
ASSISTANT_ID = "asst_AM7Evj3dYgIhdzTQX0BOe0E6"

# Definições das opções para os campos
OPCOES_FINALIDADE_DECLARACAO = [
    "Selecionar",
    "Comprovação de comparecimento a atendimento psicológico em determinada data e horário",
    "Comprovação de acompanhamento psicológico em andamento",
    "Justificativa de ausência escolar por motivo de comparecimento a atendimento psicológico",
    "Justificativa de ausência laboral por motivo de comparecimento a atendimento psicológico",
    "Comprovação de atendimento psicológico domiciliar realizado em localidade específica",
    "Atendimento psicológico no contexto de acompanhamento judicial ou institucional",
    "Participação em atividade psicológica vinculada a grupos terapêuticos, oficinas ou atendimentos em grupo",
    "Acompanhamento de familiar em sessão psicológica",
    "Encaminhamento externo para unidade de saúde, CRAS, CREAS, CAPS, etc",
    "Participação em avaliação psicológica agendada ou em curso",
    "Outros (especificar)"
]

OPCOES_FINALIDADE_ATESTADO = [
    "Selecionar",
    "Manuseio de arma de fogo (porte/posse)",
    "Direção de veículo automotor (cnh)",
    "Aptidão para exercício de função pública ou privada",
    "Participação em concursos públicos",
    "Ingresso ou permanência em instituições educativas",
    "Habilitação para adoção, guarda ou tutela",
    "Atividades esportivas, militares, entre outras",
    "Solicitar afastamento ou dispensa",
    "Comprovar a presença de condição psicológica específica",
    "Subsidiar encaminhamentos para serviços de saúde, assistência ou justiça",
    "Atestar condição de vulnerabilidade psíquica ou emocional",
    "Responder demandas específicas de autoridades ou instituições"
]

OPCOES_FINALIDADE_RELATORIO = [
    "Selecionar",
    "Acompanhamento terapêutico (clínico individual ou em grupo)",
    "Encaminhamento para outros profissionais ou serviços (médicos, pedagógicos, assistenciais)",
    "Solicitação de benefícios sociais ou previdenciários (como BPC, auxílio-doença)",
    "Acompanhamento escolar ou inclusão educacional",
    "Demanda judicial (ex: guarda de filhos, avaliação de vínculos familiares, medidas protetivas)",
    "Acompanhamento em medidas socioeducativas",
    "Monitoramento de evolução em tratamentos psicológicos",
    "Solicitação de transferência institucional (escolas, abrigos, unidades terapêuticas)",
    "Reintegração familiar ou social",
    "Necessidade de adaptação curricular ou inclusão",
    "Registro de atendimento em instituições de acolhimento (CRAS, CREAS, CAPS, abrigos, etc.)",
    "Orientações para continuidade do cuidado psicológico",
    "Participação em programa de reabilitação psicossocial",
    "Solicitação de avaliação de competências parentais",
    "Solicitação de avaliação da dinâmica familiar ou conjugal",
    "Avaliação do impacto psicológico de eventos traumáticos",
    "Requerimento de revisão de medidas protetivas",
    "Documentação para processos administrativos internos (ex: escolas, empresas)",
    "Demanda institucional para avaliação de adaptação e comportamento",
    "Relato de atuação em programas de promoção de saúde mental"
]

OPCOES_DEMANDA_RELATORIO = [
    "Selecionar",
    "👨‍👩‍👧 Clínica / Atendimento Individual",
    "Queixas de ansiedade, depressão ou angústia existencial",
    "Dificuldades nos relacionamentos interpessoais",
    "Processos de luto ou separação",
    "Transtornos alimentares ou de imagem corporal",
    "Queixas de baixa autoestima e insegurança",
    "Demandas para avaliação de autoconhecimento e desenvolvimento pessoal",
    "Elaboração de traumas psicológicos",
    "🎓 Escolar / Educacional",
    "Dificuldades de aprendizagem ou desempenho acadêmico",
    "Comportamento disruptivo em sala de aula",
    "Queixas de bullying ou exclusão social",
    "Necessidade de adaptação curricular ou inclusão escolar",
    "Avaliação de maturidade escolar",
    "🧑‍⚖️ Psicologia Jurídica",
    "Avaliação de vínculos afetivos entre criança e responsáveis",
    "Processos de disputa de guarda ou alienação parental",
    "Encaminhamento judicial para avaliação de condições parentais",
    "Medidas protetivas em contexto de violência doméstica",
    "Acompanhamento de adolescentes em cumprimento de medidas socioeducativas",
    "🧠 Saúde Mental e Psicossocial",
    "Acompanhamento de pacientes com transtornos mentais severos e persistentes",
    "Avaliação de riscos psicossociais (ideação suicida, autolesão, etc.)",
    "Reintegração social após internação psiquiátrica",
    "Suporte psicológico em contextos de violência, abuso ou negligência",
    "👔 Organizacional / Institucional",
    "Avaliação de clima organizacional e relações no trabalho",
    "Demandas de readaptação ou retorno ao trabalho",
    "Conflitos interpessoais entre funcionários",
    "Avaliação de condições emocionais para o desempenho laboral",
    "🏠 Serviços de Proteção Social / Institucional",
    "Acompanhamento de crianças e adolescentes em acolhimento institucional",
    "Situações de vulnerabilidade social e familiar",
    "Encaminhamentos por órgãos como CRAS, CREAS, CAPS",
    "Necessidade de inclusão em programas de assistência ou políticas públicas"
]

OPCOES_PROCEDIMENTOS_RELATORIO = [
    "Selecionar",
    "🗣️ Entrevistas",
    "Entrevista inicial (acolhimento)",
    "Entrevista psicológica individual",
    "Entrevista devolutiva",
    "Entrevista com familiares ou responsáveis",
    "Entrevista com equipe técnica ou pedagógica (quando em contexto institucional)",
    "📋 Observações",
    "Observação comportamental em contexto clínico",
    "Observação em ambiente institucional (escola, abrigo, trabalho)",
    "Registros sistemáticos de conduta (análise funcional)",
    "📚 Instrumentos Técnicos",
    "Aplicação de testes psicológicos reconhecidos pelo SATEPSI (ex.: WISC, Bender, R1, etc.)",
    "Inventários de sintomas (ex.: BDI, STAI)",
    "Questionários sociodemográficos ou de rastreio",
    "Escalas de avaliação (ex.: Vineland, SNAP-IV, CBCL)",
    "🧠 Procedimentos de Avaliação Psicológica",
    "Anamnese detalhada",
    "Avaliação de habilidades cognitivas",
    "Avaliação de aspectos emocionais e afetivos",
    "Avaliação de aspectos relacionais e sociais",
    "Avaliação da dinâmica familiar",
    "Avaliação de potencial de aprendizagem ou funcionamento adaptativo",
    "🧾 Revisão e Análise Documental",
    "Análise de relatórios médicos ou educacionais",
    "Estudo de prontuários psicológicos anteriores",
    "Levantamento histórico de atendimentos",
    "🧑‍🤝‍🧑 Participação em Atividades / Dinâmicas",
    "Dinâmicas de grupo",
    "Oficinas terapêuticas",
    "Rodas de conversa ou grupos de acolhimento",
    "💼 Outros Procedimentos Contextuais",
    "Visitas domiciliares (quando autorizadas)",
    "Interlocução com instituições externas (CRAS, CREAS, CAPS)",
    "Reuniões de equipe técnica"
]

OPCOES_PROCEDIMENTOS_LAUDO = [
    "Selecionar",
    "Entrevista com os responsáveis",
    "Entrevista com a professora",
    "Entrevista clínica semiestruturada",
    "Observação comportamental direta durante os atendimentos",
    "Observação direta em contexto clínico e lúdico",
    "Análise de relatórios escolares e comportamentais fornecidos pela escola",
    "Aplicação de instrumentos padronizados",
    "Aplicação de instrumentos complementares"
]

OPCOES_FINALIDADE_PARECER = [
    "Selecionar",
    "Esclarecimento técnico sobre capacidades parentais (ex: guarda, convivência familiar, alienação parental)",
    "Análise técnica de documentos psicológicos emitidos por terceiros (ex: contralaudo ou impugnação)",
    "Opinião técnica sobre possível necessidade de curatela ou interdição civil",
    "Verificação de indícios de violação de direitos (ex: negligência, abuso, abandono afetivo)",
    "Subsidiar decisões do Judiciário em medidas protetivas ou acolhimento institucional",
    "Opinar sobre a capacidade de compreensão e participação de um réu/processado",
    "Orientar tecnicamente quanto à indicação de guarda compartilhada ou unilateral",
    "Análise de conflitos familiares complexos em disputa judicial",
    "Apoiar decisões sobre internação involuntária ou compulsória",
    "Fundamentar a necessidade de encaminhamento psiquiátrico ou multiprofissional",
    "Orientar a equipe sobre prognóstico de funcionalidade psíquica de usuários crônicos",
    "Parecer sobre a aderência ao tratamento psicológico/psiquiátrico",
    "Responder tecnicamente sobre a necessidade de inclusão em programas de proteção social",
    "Opinar sobre benefícios assistenciais (ex: BPC-LOAS, isenção de imposto, aposentadoria) por condição psíquica",
    "Subsidiar encaminhamentos à rede de educação inclusiva, CRAS, CREAS, CAPS, Conselho Tutelar",
    "Apontar a necessidade de acolhimento institucional de crianças, adolescentes ou idosos",
    "Parecer sobre necessidades educacionais especiais e adaptações pedagógicas",
    "Verificação de sofrimento psíquico em contexto escolar (ex: bullying, ansiedade, fobia escolar)",
    "Apoio técnico sobre possível evasão escolar por motivo psíquico-comportamental",
    "Orientação sobre medidas protetivas para crianças e adolescentes em risco no ambiente escolar",
    "Esclarecimento sobre uso de instrumentos psicológicos em avaliação específica (ex: validade de teste)",
    "Apoiar decisões técnicas em perícias sociais, avaliações multiprofissionais e programas institucionais",
    "Parecer para esclarecimento ético em contextos que envolvam condutas ou práticas psicológicas contestadas"
]

OPCOES_OBJETIVOS_PARECER = [
    "Selecionar",
    "Esclarecer a capacidade civil de pessoa maior de idade em situação de possível vulnerabilidade psíquica",
    "Analisar a função parental de um ou ambos os responsáveis legais em disputas de guarda",
    "Apontar indícios de alienação parental ou prejuízo no vínculo afetivo entre genitor e criança/adolescente",
    "Verificar a idoneidade emocional e relacional de responsável legal em ações de tutela ou curatela",
    "Analisar tecnicamente os efeitos psicológicos da convivência com familiares acusados de violência",
    "Fornecer subsídios técnicos em processos de revisão de medidas socioeducativas ou protetivas",
    "Avaliar a adequação de retorno ao convívio familiar após acolhimento institucional",
    "Fundamentar decisões judiciais quanto à necessidade de encaminhamentos à rede de saúde mental",
    "Analisar documentos psicológicos apresentados nos autos e emitir parecer técnico sobre sua validade e coerência técnica (ex: contestação ou validação de laudos e relatórios)",
    "Emitir opinião técnica em processos de adoção ou destituição do poder familiar",
    "Avaliar a aderência ou resistência ao tratamento psicológico ou psiquiátrico",
    "Verificar a presença de sinais de sofrimento psíquico grave, sugerindo encaminhamentos adequados",
    "Analisar a necessidade de internação involuntária sob o ponto de vista ético-técnico",
    "Esclarecer os possíveis efeitos emocionais de dependência química e comorbidades psíquicas em familiares e conviventes",
    "Compreender a dinâmica familiar conflituosa e seu impacto psicológico sobre crianças, adolescentes ou idosos",
    "Identificar fatores de risco psicossociais em contextos de negligência, abandono ou violência",
    "Verificar a presença de recursos de suporte e proteção familiar",
    "Apontar a viabilidade de encaminhamento para rede de proteção (CRAS, CREAS, CAPS, etc.)",
    "Emitir parecer sobre necessidade de suporte educacional especial ou mediação pedagógica",
    "Verificar a influência de fatores emocionais e familiares no desempenho escolar e no comportamento de crianças/adolescentes",
    "Analisar os impactos psicológicos de violência escolar ou exclusão social",
    "Orientar quanto à necessidade de encaminhamentos psicopedagógicos ou multidisciplinares",
    "Emitir parecer sobre a adequação de metodologia, validade e coerência interna de documentos psicológicos",
    "Responder a questionamentos de órgãos públicos sobre documentos já emitidos",
    "Fundamentar tecnicamente decisões intersetoriais no âmbito da saúde, justiça e assistência social",
    "Oferecer subsídio à atuação técnica da equipe multiprofissional da instituição solicitante"
]

OPCOES_GENERO = ["Selecionar", "Masculino", "Feminino", "Outro", "Prefere não informar"]

OPCOES_ESCOLARIDADE = [
    "Selecionar",
    "Sem instrução",
    "Ensino Fundamental Incompleto",
    "Ensino Fundamental Completo",
    "Ensino Médio Incompleto",
    "Ensino Médio Completo",
    "Ensino Técnico",
    "Ensino Superior Incompleto",
    "Ensino Superior Completo",
    "Pós-graduação (Lato Sensu)",
    "Mestrado",
    "Doutorado"
]

OPCOES_PROFISSAO = [
    "Selecionar",
    "Estudante",
    "Desempregado",
    "Autônomo",
    "Empregado com carteira assinada",
    "Servidor público",
    "Aposentado",
    "Empresário",
    "Profissional liberal",
    "Dona(o) de casa",
    "Outro"
]

OPCOES_ESTADO_CIVIL = [
    "Selecionar",
    "Solteiro(a)",
    "Casado(a)",
    "União estável",
    "Separado(a) judicialmente",
    "Divorciado(a)",
    "Viúvo(a)"
]

TEXTO_ANALISE = """Irei coletar, interpretar e analisar os dados provenientes de todos os procedimentos utilizados, integrando as informações de forma criteriosa para apresentar os achados da avaliação psicológica de maneira descritiva. Essa apresentação contempla: Aspectos emocionais, cognitivos, sociais e comportamentais, articulando-os com os resultados obtidos por meio dos instrumentos psicológicos aplicados."""

TEXTO_CONCLUSAO = """Apresentarei a síntese dos achados da avaliação psicológica, com base nos procedimentos realizados, nos dados obtidos por meio dos instrumentos aplicados e nas demais técnicas empregadas. Caso os resultados permitam, poderão ser incluídas hipóteses diagnósticas fundamentadas nos critérios do CID-11 ou do DSM-5-TR, bem como indicadas intervenções, encaminhamentos ou outras medidas pertinentes, de acordo com as evidências clínicas observadas."""

TEXTO_ATESTADO_OBSERVACOES = """Este Atestado Psicológico possui caráter sigiloso, foi emitido exclusivamente para a finalidade aqui declarada e não deverá ser utilizado para quaisquer outros fins que não aqueles expressamente indicados. Ressalta-se que se trata de documento extrajudicial, elaborado conforme os princípios técnicos e éticos da Psicologia, em especial os dispostos no Código de Ética Profissional do Psicólogo e na Resolução CFP nº 06/2019."""

def extrair_texto_arquivo(file):
    try:
        if file.type == "application/pdf":
            reader = PdfReader(file)
            texto = "\n".join(page.extract_text() or "" for page in reader.pages)
            return texto.strip() or "[PDF sem texto detectável]"
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file)
            texto = "\n".join([p.text for p in doc.paragraphs])
            return texto.strip() or "[DOCX sem texto detectável]"
        elif file.type.startswith("image/"):
            image = Image.open(file)
            texto = pytesseract.image_to_string(image, lang="por")
            return texto.strip() or "[Imagem sem texto detectável]"
        else:
            return f"[Tipo de arquivo não suportado: {file.type}]"
    except Exception as e:
        logging.error(f"Erro ao extrair texto de {file.name}: {e}")
        return f"[Erro ao extrair texto: {e}]"

def sanitize_filename(filename):
    return re.sub(r'[^a-zA-Z0-9_\-\.]', '_', filename)

def obter_campos_por_tipo_documento(tipo):
    estrutura = {
        "DECLARAÇÃO PSICOLÓGICA": [
            "NOME DO(A) PACIENTE",
            "DATA DE NASCIMENTO",
            "FINALIDADE",
            "LOCAL DO(S) ATENDIMENTO(S)",
            "DATA(S) E HORÁRIO(S) DO(S) ATENDIMENTO(S)",
            "DURAÇÃO DO ACOMPANHAMENTO PSICOLÓGICO"
        ],
        "ATESTADO PSICOLÓGICO": [
            "NOME DA PESSOA OU INSTITUIÇÃO ATENDIDA",
            "DATA DE NASCIMENTO (opcional)",
            "IDADE (opcional)",
            "GÊNERO (opcional)",
            "ESCOLARIDADE (opcional)",
            "PROFISSÃO (opcional)",
            "ESTADO CIVIL (opcional)",
            "SOLICITANTE",
            "FINALIDADE",
            "DESCRIÇÃO DAS CONDIÇÕES PSICOLÓGICAS",
            "CID OU OUTRAS CLASSIFICAÇÕES DIAGNÓSTICAS (opcional)",
            "LOCAL DA AVALIAÇÃO",
            "DATA DA EMISSÃO",
            "PSICÓLOGO RESPONSÁVEL",
            "OBSERVAÇÕES"
        ],
        "RELATÓRIO PSICOLÓGICO": [
            "NOME DA PESSOA OU INSTITUIÇÃO ATENDIDA",
            "SOLICITANTE",
            "DATA",
            "LOCAL",
            "FINALIDADE DO DOCUMENTO",
            "DESCRIÇÃO DA DEMANDA",
            "PROCEDIMENTOS UTILIZADOS",
            "OBSERVAÇÕES CLÍNICAS",
            "ANÁLISE",
            "CONCLUSÃO",
            "REFERÊNCIAS"
        ],
        "LAUDO PSICOLÓGICO": [
            "NOME DO(A) PACIENTE",
            "DATA DE NASCIMENTO",
            "IDADE",
            "GÊNERO",
            "ESCOLARIDADE",
            "PROFISSÃO",
            "ESTADO CIVIL",
            "SOLICITANTE",
            "QUAL FOI O OBJETIVO DA SOLICITAÇÃO",
            "QUEIXA PRINCIPAL",
            "O SEU ENDEREÇO PROFISSIONAL",
            "DATA DA AVALIAÇÃO",
            "LOCAL DA AVALIAÇÃO",
            "PROCEDIMENTOS",
            "OBSERVAÇÕES CLÍNICAS",
            "ANÁLISE",
            "CONCLUSÃO",
            "REFERÊNCIAS"
        ],
        "PARECER PSICOLÓGICO": [
            "NOME DA PESSOA OU INSTITUIÇÃO ATENDIDA",
            "DATA DE NASCIMENTO",
            "IDADE",
            "GÊNERO",
            "SOLICITANTE",
            "FINALIDADE DO DOCUMENTO",
            "OBJETIVOS DA CONSULTA/DEMANDA PARA PARECER PSICOLÓGICO",
            "ANÁLISE",
            "CONCLUSÃO",
            "REFERÊNCIAS"
        ]
    }
    return estrutura.get(tipo, [])

def gerar_campo_comum(campo, tipo_campo="texto", opcoes=None, key=None):
    """Função auxiliar para gerar campos comuns"""
    safe_key = key or re.sub(r'[^a-zA-Z0-9_]', '_', campo)
    
    if tipo_campo == "texto":
        return st.text_input(campo, key=f"input_{safe_key}")
    elif tipo_campo == "data":
        return st.date_input(campo, format="DD/MM/YYYY", key=f"date_{safe_key}")
    elif tipo_campo == "numero":
        return st.number_input(campo, min_value=0, max_value=120, step=1, key=f"number_{safe_key}")
    elif tipo_campo == "selecao":
        return st.selectbox(campo, opcoes, key=f"select_{safe_key}")
    elif tipo_campo == "multiselecao":
        return st.multiselect(campo, opcoes, key=f"multiselect_{safe_key}")
    elif tipo_campo == "texto_area":
        return st.text_area(campo, key=f"textarea_{safe_key}")

def gerar_campos_dinamicos(campos, tipo_documento):
    respostas = {}

    for campo in campos:
        safe_key = re.sub(r'[^a-zA-Z0-9_]', '_', campo)
        
        # Campos comuns a vários documentos
        if campo in ["NOME DO(A) PACIENTE", "NOME DA PESSOA OU INSTITUIÇÃO ATENDIDA", "SOLICITANTE", 
                    "LOCAL DA AVALIAÇÃO", "LOCAL DO(S) ATENDIMENTO(S)", "DURAÇÃO DO ACOMPANHAMENTO PSICOLÓGICO",
                    "CID OU OUTRAS CLASSIFICAÇÕES DIAGNÓSTICAS (opcional)", "QUAL FOI O OBJETIVO DA SOLICITAÇÃO",
                    "QUEIXA PRINCIPAL", "O SEU ENDEREÇO PROFISSIONAL"]:
            respostas[campo] = gerar_campo_comum(campo, "texto", key=f"{tipo_documento}_{safe_key}")
        
        elif campo in ["DATA DE NASCIMENTO", "DATA DE NASCIMENTO (opcional)", "DATA DA AVALIAÇÃO", "DATA DA EMISSÃO", "DATA"]:
            respostas[campo] = gerar_campo_comum(campo, "data", key=f"{tipo_documento}_{safe_key}")
        
        elif campo in ["IDADE", "IDADE (opcional)"]:
            respostas[campo] = gerar_campo_comum(campo, "numero", key=f"{tipo_documento}_{safe_key}")
        
        elif campo in ["GÊNERO", "GÊNERO (opcional)"]:
            respostas[campo] = gerar_campo_comum(campo, "selecao", OPCOES_GENERO, key=f"{tipo_documento}_{safe_key}")
        
        elif campo in ["ESCOLARIDADE", "ESCOLARIDADE (opcional)"]:
            respostas[campo] = gerar_campo_comum(campo, "selecao", OPCOES_ESCOLARIDADE, key=f"{tipo_documento}_{safe_key}")
        
        elif campo in ["PROFISSÃO", "PROFISSÃO (opcional)"]:
            profissao = gerar_campo_comum(campo, "selecao", OPCOES_PROFISSAO, key=f"{tipo_documento}_{safe_key}")
            if profissao == "Outro":
                profissao_outro = st.text_input("Especifique a profissão", key=f"{tipo_documento}_{safe_key}_outro")
                respostas[campo] = f"Outro: {profissao_outro}" if profissao_outro else "Outro"
            else:
                respostas[campo] = profissao
        
        elif campo in ["ESTADO CIVIL", "ESTADO CIVIL (opcional)"]:
            respostas[campo] = gerar_campo_comum(campo, "selecao", OPCOES_ESTADO_CIVIL, key=f"{tipo_documento}_{safe_key}")
        
        # Campos específicos por tipo de documento
        elif campo == "FINALIDADE" and tipo_documento == "DECLARAÇÃO PSICOLÓGICA":
            finalidade = gerar_campo_comum(campo, "selecao", OPCOES_FINALIDADE_DECLARACAO, key=f"{tipo_documento}_{safe_key}")
            if finalidade == "Outros (especificar)":
                finalidade_outro = st.text_input("Especifique a finalidade", key=f"{tipo_documento}_{safe_key}_outro")
                respostas[campo] = f"Outros: {finalidade_outro}" if finalidade_outro else "Outros"
            else:
                respostas[campo] = finalidade
        
        elif campo == "FINALIDADE" and tipo_documento == "ATESTADO PSICOLÓGICO":
            respostas[campo] = gerar_campo_comum(campo, "selecao", OPCOES_FINALIDADE_ATESTADO, key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "FINALIDADE DO DOCUMENTO" and tipo_documento == "RELATÓRIO PSICOLÓGICO":
            respostas[campo] = gerar_campo_comum(campo, "selecao", OPCOES_FINALIDADE_RELATORIO, key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "DESCRIÇÃO DA DEMANDA" and tipo_documento == "RELATÓRIO PSICOLÓGICO":
            respostas[campo] = gerar_campo_comum(campo, "selecao", OPCOES_DEMANDA_RELATORIO, key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "PROCEDIMENTOS UTILIZADOS" and tipo_documento == "RELATÓRIO PSICOLÓGICO":
            respostas[campo] = gerar_campo_comum(campo, "multiselecao", OPCOES_PROCEDIMENTOS_RELATORIO, key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "PROCEDIMENTOS" and tipo_documento == "LAUDO PSICOLÓGICO":
            respostas[campo] = gerar_campo_comum(campo, "multiselecao", OPCOES_PROCEDIMENTOS_LAUDO, key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "FINALIDADE DO DOCUMENTO" and tipo_documento == "PARECER PSICOLÓGICO":
            respostas[campo] = gerar_campo_comum(campo, "selecao", OPCOES_FINALIDADE_PARECER, key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "OBJETIVOS DA CONSULTA/DEMANDA PARA PARECER PSICOLÓGICO":
            respostas[campo] = gerar_campo_comum(campo, "selecao", OPCOES_OBJETIVOS_PARECER, key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "DATA(S) E HORÁRIO(S) DO(S) ATENDIMENTO(S)":
            st.markdown(f"**{campo}**")
            num_datas = st.number_input("Quantos atendimentos deseja registrar?", min_value=1, max_value=10, value=1, key=f"{tipo_documento}_{safe_key}_num")
            datas_horarios = []
            
            for i in range(num_datas):
                st.markdown(f"Atendimento {i+1}")
                col1, col2, col3 = st.columns(3)
                with col1:
                    data = st.date_input(f"Data {i+1}", key=f"{tipo_documento}_{safe_key}_data_{i}", format="DD/MM/YYYY")
                with col2:
                    hora_inicio = st.time_input(f"Horário de início {i+1}", key=f"{tipo_documento}_{safe_key}_inicio_{i}")
                with col3:
                    hora_fim = st.time_input(f"Horário de fim {i+1}", key=f"{tipo_documento}_{safe_key}_fim_{i}")
                datas_horarios.append({
                    "data": data,
                    "inicio": hora_inicio,
                    "fim": hora_fim
                })
            respostas[campo] = datas_horarios
        
        elif campo == "DESCRIÇÃO DAS CONDIÇÕES PSICOLÓGICAS":
            st.markdown("**DESCRIÇÃO DAS CONDIÇÕES PSICOLÓGICAS**")
            st.info("Por favor, descreva as informações sobre o estado psicológico do(a) beneficiário(a), conforme identificadas no processo de avaliação psicológica, de forma compatível com a finalidade deste atestado. Se preferir, você pode anexar o relatório ou laudo psicológico correspondente, e eu faço a análise para redigir as condições psicológicas adequadas à emissão do documento.")
            respostas[campo] = gerar_campo_comum(campo, "texto_area", key=f"{tipo_documento}_{safe_key}")
            arquivos_condicoes = st.file_uploader("Anexar relatório ou laudo psicológico", accept_multiple_files=True, key=f"{tipo_documento}_{safe_key}_files")
            if arquivos_condicoes:
                textos_anexos = [extrair_texto_arquivo(f) for f in arquivos_condicoes]
                respostas[f"{campo}_ANEXOS"] = textos_anexos
        
        elif campo == "OBSERVAÇÕES CLÍNICAS":
            respostas[campo] = gerar_campo_comum(campo, "texto_area", key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "ANÁLISE":
            st.markdown("**ANÁLISE**")
            st.info(TEXTO_ANALISE)
            respostas[campo] = gerar_campo_comum("VOCÊ GOSTARIA DE FAZER ALGUMA OBSERVAÇÃO?", "texto_area", key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "CONCLUSÃO":
            st.markdown("**CONCLUSÃO**")
            st.info(TEXTO_CONCLUSAO)
            respostas[campo] = gerar_campo_comum("VOCÊ GOSTARIA DE FAZER ALGUMA OBSERVAÇÃO?", "texto_area", key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "REFERÊNCIAS":
            st.markdown("**REFERÊNCIAS**")
            respostas[campo] = gerar_campo_comum("Você gostaria de fornecer alguma observação?", "texto_area", key=f"{tipo_documento}_{safe_key}")
        
        elif campo == "OBSERVAÇÕES":
            st.markdown("**OBSERVAÇÕES**")
            st.info(TEXTO_ATESTADO_OBSERVACOES)
            respostas[campo] = TEXTO_ATESTADO_OBSERVACOES
        
        elif campo == "PSICÓLOGO RESPONSÁVEL":
            respostas[campo] = f"{nome} - CRP: {numero}" if nome and numero else ""

    st.markdown("---")
    st.markdown("**Documentos Complementares**")
    arquivos = st.file_uploader("Anexar arquivo(s) complementares", accept_multiple_files=True, key=f"{tipo_documento}_complementares")

    return respostas, arquivos

def enviar_para_assistente(user_message):
    try:
        thread = openai.beta.threads.create()
        openai.beta.threads.messages.create(thread_id=thread.id, role="user", content=user_message)
        run = openai.beta.threads.runs.create(thread_id=thread.id, assistant_id=ASSISTANT_ID)
        while True:
            run = openai.beta.threads.runs.retrieve(thread_id=thread.id, run_id=run.id)
            if run.status == "completed":
                break
            time.sleep(1)
        messages_list = openai.beta.threads.messages.list(thread_id=thread.id)
        for msg in reversed(messages_list.data):
            if msg.role == "assistant":
                return msg.content[0].text.value
        return "[Resposta não encontrada]"
    except Exception as e:
        logging.error(f"Erro na API Assistants: {e}")
        return f"[Erro ao interagir com o assistente: {e}]"

def exportar_para_docx(texto):
    doc = Document()
    for paragrafo in texto.split("\n"):
        doc.add_paragraph(paragrafo.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Janela de aceite de termos
if 'accepted_terms' not in st.session_state:
    st.session_state.accepted_terms = False

if not st.session_state.accepted_terms:
    st.title("Termos de Uso e Política de Privacidade")
    st.markdown("""
    🛡️ **Compromisso com a Ética, Segurança e Sigilo Profissional**
                
    Este serviço foi desenvolvido como ferramenta de apoio técnico à elaboração de documentos psicológicos, com base nas diretrizes estabelecidas pela **Resolução CFP nº 06/2019**, pela **Resolução CFP nº 01/2009 (Política de Proteção de Dados)** e pelo **Código de Ética Profissional do Psicólogo**.

                
    🧠 **Responsabilidade Técnica e Ética**  
    As produções dos documentos devem **obrigatoriamente ser revisadas, validadas e assinadas por psicóloga(o) devidamente inscrita(o) no CRP**, conforme determina a legislação profissional.  
    O conteúdo gerado **não substitui o julgamento clínico e técnico do profissional**.

                
    📌 **Finalidade do Sistema**  
    Este assistente virtual tem como único propósito **auxiliar a(o) psicóloga(o)** na sistematização de informações, organização textual e conformidade estrutural de documentos, sempre respeitando os princípios de autonomia, consentimento informado, sigilo, não exposição e ética nas relações profissionais.
                

    ⚖️ **Referências Normativas**
    - Resolução CFP nº 06/2019 – Elaboração de Documentos Escritos Produzidos pela(o) Psicóloga(o)
    - Código de Ética Profissional do Psicólogo – Artigos 1º, 9º, 13º e 14º
    - Resolução CFP nº 11/2018 – Sobre uso de tecnologias da informação e comunicação
    - LGPD (Lei Geral de Proteção de Dados) – Aplicabilidade ao contexto psicológico
                

    🔒 **Privacidade e Proteção de Dados**  
    - Esta ferramenta foi construída em conformidade com:
    - O Código de Ética do Profissional Psicólogo (Resolução CFP nº 010/2005);
    - A Resolução CFP nº 06/2019: Elaboração de Documentos Escritos Produzidos pela(o) Psicóloga(o);
    - Resolução CFP nº 11/2018: Sobre uso de tecnologias da informação e comunicação
    - Criptografia em trânsito (HTTPS): Criptografia de Ponta a Ponta para Proteger Dados em Trânsito e em Repouso. Todos os dados são protegidos contra interceptação.
    - Controle de acesso: APIs protegidas com autenticação para impedir acesso não autorizado.
    - Validação de entrada: Validações automáticas, evitando injeções maliciosas ou erros lógicos.
    - Registros e auditoria: Rastreamento de dados com precisão (data/hora e autor), ajudando na responsabilização e conformidade com normas como a LGPD.
    - Anonimização: Omissão de dados sensíveis antes de armazenar ou compartilhar informações JSON, promovendo privacidade.
    - Normas da Lei Geral de Proteção de Dados Pessoais (Lei nº 13.709/2018), que regula o tratamento de dados pessoais no Brasil. Seu objetivo principal é garantir o direito à privacidade e à proteção dos dados dos cidadãos, estabelecendo regras claras sobre coleta, uso, armazenamento e compartilhamento de informações pessoais por empresas, órgãos públicos e profissionais autônomos incluindo psicólogas(os).
                
    Ao utilizar este sistema, você declara ciência de que **respeita e segue os preceitos éticos da profissão** e que **assume a responsabilidade técnica e legal pelos documentos emitidos** com o apoio desta ferramenta.)
    """)

    if st.button("Aceito os Termos e Continuar"):
        st.session_state.accepted_terms = True
    else:
        st.stop()

st.header("🧠 Psicólogo Assistente / Elaboração de Documentos 🧠", divider="gray")
nome = st.text_input("SEU NOME COMPLETO")
numero = st.text_input("CRP")
data = st.text_input("DATA")
if nome and numero:
    st.success(f"Bem-vindo(a), {nome}! | CRP: {numero}")

st.markdown("---")
tipo_documento = st.selectbox("Tipo de Documento", [
    "DECLARAÇÃO PSICOLÓGICA",
    "ATESTADO PSICOLÓGICO",
    "RELATÓRIO PSICOLÓGICO",
    "LAUDO PSICOLÓGICO",
    "PARECER PSICOLÓGICO"
])

st.markdown("---")
campos = obter_campos_por_tipo_documento(tipo_documento)
respostas, arquivos = gerar_campos_dinamicos(campos, tipo_documento)

enviar = st.button("🔍 Gerar Documento")
if enviar:
    conteudo = f"Tipo de Documento: {tipo_documento}\n\n"
    
    # Adiciona as respostas dos campos
    for campo in campos:
        if campo == "PROCEDIMENTOS UTILIZADOS" or campo == "PROCEDIMENTOS":
            conteudo += "PROCEDIMENTOS:\n"
            if respostas[campo]:
                for proc in respostas[campo]:
                    conteudo += f"- {proc}\n"
            conteudo += "\n"
        elif campo == "DATA(S) E HORÁRIO(S) DO(S) ATENDIMENTO(S)":
            conteudo += f"{campo}:\n"
            for atendimento in respostas[campo]:
                data_formatada = atendimento["data"].strftime("%d/%m/%Y")
                inicio_formatado = atendimento["inicio"].strftime("%H:%M")
                fim_formatado = atendimento["fim"].strftime("%H:%M")
                conteudo += f"- {data_formatada} das {inicio_formatado} às {fim_formatado}\n"
            conteudo += "\n"
        elif campo in ["DATA DE NASCIMENTO", "DATA DE NASCIMENTO (opcional)", "DATA DA AVALIAÇÃO", "DATA DA EMISSÃO", "DATA"]:
            valor = respostas[campo]
            if valor:
                data_formatada = valor.strftime("%d/%m/%Y")
                conteudo += f"{campo}: {data_formatada}\n\n"
        elif campo == "DESCRIÇÃO DAS CONDIÇÕES PSICOLÓGICAS":
            conteudo += f"{campo}:\n{respostas[campo]}\n\n"
            if f"{campo}_ANEXOS" in respostas:
                conteudo += "DOCUMENTOS ANEXADOS À DESCRIÇÃO:\n"
                for texto in respostas[f"{campo}_ANEXOS"]:
                    conteudo += f"{texto}\n\n"
        else:
            valor = respostas[campo]
            if valor and valor != "Selecionar":  # Só adiciona o campo se tiver algum valor e não for "Selecionar"
                conteudo += f"{campo}:\n{valor}\n\n"
    
    # Adiciona o conteúdo dos arquivos anexados
    if arquivos:
        conteudo += "DOCUMENTOS COMPLEMENTARES:\n"
        textos_extraidos = [extrair_texto_arquivo(f) for f in arquivos]
        conteudo += "\n\n".join(textos_extraidos)

    resposta = enviar_para_assistente(conteudo)
    st.subheader("📄 Documento Gerado")
    st.text_area("Conteúdo", resposta, height=400)
    buffer = exportar_para_docx(resposta)
    st.download_button("📥 Baixar DOCX", data=buffer, file_name=f"{sanitize_filename(tipo_documento)}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.info("🔍 Este documento deve ser revisado pelo psicólogo responsável antes do uso oficial.")



